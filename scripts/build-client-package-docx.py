#!/usr/bin/env python3
# Generates public/ppwr-klienta-paketes-sablons.docx from the canonical
# Markdown source (public/ppwr-klienta-paketes-sablons.md). The Markdown file
# is the single source of truth — edit it, then re-run this script; never edit
# the .docx by hand.
#
# Supported Markdown subset: # / ## / ### headings, the paragraph immediately
# after the H1 title (rendered as a subtitle), > blockquotes, --- rules,
# GFM tables, paragraphs with **bold**, _italic_ and `code` inline markup.
#
# Requires python-docx:  pip install python-docx
# Run from repo root:     python3 scripts/build-client-package-docx.py
import os
import re

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor

SRC = os.path.join(os.path.dirname(__file__), "..", "public", "ppwr-klienta-paketes-sablons.md")
OUT = os.path.join(os.path.dirname(__file__), "..", "public", "ppwr-klienta-paketes-sablons.docx")

GREY = RGBColor(0x80, 0x80, 0x80)
DARK_GREY = RGBColor(0x55, 0x55, 0x55)
HEADER_FILL = "E6E6E6"
NOTE_FILL = "F2F2F2"

INLINE_RE = re.compile(r"(\*\*.+?\*\*|_[^_]+_|`[^`]+`)")


def add_runs(paragraph, text, size=None, color=None, italic=False):
    """Render inline **bold** / _italic_ / `code` markup into runs."""
    for token in INLINE_RE.split(text):
        if not token:
            continue
        bold = False
        it = italic
        if token.startswith("**") and token.endswith("**") and len(token) > 4:
            token, bold = token[2:-2], True
        elif token.startswith("_") and token.endswith("_") and len(token) > 2:
            token, it = token[1:-1], True
        elif token.startswith("`") and token.endswith("`") and len(token) > 2:
            token = token[1:-1]  # backticks are Markdown-only; render as plain text
        run = paragraph.add_run(token)
        run.bold = bold
        run.italic = it
        if size:
            run.font.size = size
        if color:
            run.font.color.rgb = color


def set_cell_bg(cell, hex_fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_fill)
    tc_pr.append(shd)


def shade_paragraph(paragraph, hex_fill):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hex_fill)
    p_pr.append(shd)


def add_hr(doc):
    p = doc.add_paragraph()
    p_pr = p._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "808080")
    borders.append(bottom)
    p_pr.append(borders)


def split_row(line):
    return [c.strip() for c in line.strip().strip("|").split("|")]


def add_table(doc, headers, rows):
    font_size = Pt(9.5) if len(headers) <= 4 else Pt(8.5)
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        add_runs(hdr[i].paragraphs[0], h, size=font_size)
        for run in hdr[i].paragraphs[0].runs:
            run.bold = True
        set_cell_bg(hdr[i], HEADER_FILL)
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row[: len(headers)]):
            cells[i].text = ""
            color = GREY if val.startswith("[") else None
            add_runs(cells[i].paragraphs[0], val, size=font_size, color=color)
    if len(headers) == 2:
        for row in table.rows:
            row.cells[0].width = Pt(200)
            row.cells[1].width = Pt(240)
    return table


def main():
    with open(os.path.normpath(SRC), encoding="utf-8") as f:
        lines = f.read().splitlines()

    doc = Document()
    doc.core_properties.title = "Iepakojuma tehnisko datu veidlapa klientam"
    doc.core_properties.author = "ppwr.clickscale.dev"

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)

    saw_title = False
    subtitle_pending = False
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        if not line.strip():
            i += 1
            continue

        if line.startswith("### "):
            doc.add_heading(line[4:], level=2)
            subtitle_pending = False
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=1)
            subtitle_pending = False
        elif line.startswith("# "):
            doc.add_heading(line[2:], level=0)
            saw_title, subtitle_pending = True, True
        elif line.strip() == "---":
            add_hr(doc)
            subtitle_pending = False
        elif line.startswith("> "):
            p = doc.add_paragraph()
            add_runs(p, line[2:])
            shade_paragraph(p, NOTE_FILL)
            subtitle_pending = False
        elif line.startswith("|"):
            headers = split_row(line)
            rows = []
            i += 1
            if i < len(lines) and re.fullmatch(r"\|?[\s:|-]+\|?", lines[i].strip()):
                i += 1
            while i < len(lines) and lines[i].strip().startswith("|"):
                rows.append(split_row(lines[i]))
                i += 1
            add_table(doc, headers, rows)
            subtitle_pending = False
            continue
        else:
            if saw_title and subtitle_pending:
                p = doc.add_paragraph()
                add_runs(p, line, size=Pt(12), color=DARK_GREY, italic=True)
                subtitle_pending = False
            elif line.startswith("_") and line.endswith("_"):
                p = doc.add_paragraph()
                add_runs(p, line[1:-1], size=Pt(9), color=DARK_GREY, italic=True)
            else:
                p = doc.add_paragraph()
                add_runs(p, line)
        i += 1

    doc.save(os.path.normpath(OUT))
    print("Wrote", os.path.normpath(OUT))


if __name__ == "__main__":
    main()
